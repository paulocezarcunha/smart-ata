import streamlit as st
import google.generativeai as genai
from openai import OpenAI
from docx import Document
from io import BytesIO
import os

# Configuração visual do App
st.set_page_config(page_title="SmartAta SAP PMO", page_icon="💼")
st.title("💼 SmartAta: Governança PMO Sênior")

# 1. Configuração de Chaves (Secrets ou Input)
# Dica: No Streamlit Cloud, você pode esconder isso nos 'Secrets'
with st.sidebar:
    st.header("Configurações de Acesso")
    oa_key = st.text_input("OpenAI API Key", type="password")
    gem_key = st.text_input("Gemini API Key", type="password")

# 2. Função para criar o documento Word
def criar_docx(texto):
    doc = Document()
    doc.add_heading('ATA DE REUNIÃO EXECUTIVA', 0)
    for p in texto.split('\n'):
        doc.add_paragraph(p)
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# 3. Interface de Upload
arquivo_audio = st.file_uploader("Carregue o áudio da reunião (.mp3, .m4a, .wav)", type=['mp3', 'm4a', 'wav'])

if arquivo_audio and st.button("🚀 Gerar Ata Executiva"):
    if not oa_key or not gem_key:
        st.error("Insira as chaves de API na barra lateral!")
    else:
        try:
            # --- FASE 1: TRANSCRIÇÃO (OpenAI) ---
            client_oa = OpenAI(api_key=oa_key)
            with st.spinner("👂 Ouvindo e transcrevendo (Whisper API)..."):
                # Salva o arquivo temporariamente
                with open("temp.mp3", "wb") as f:
                    f.write(arquivo_audio.read())
                
                with open("temp.mp3", "rb") as audio_data:
                    transcricao = client_oa.audio.transcriptions.create(
                        model="whisper-1", 
                        file=audio_data
                    ).text

            # --- FASE 2: INTELIGÊNCIA PMO (Gemini) ---
            genai.configure(api_key=gem_key)
            model = genai.GenerativeModel('gemini-1.5-flash')
            
            with st.spinner("🧠 Estruturando governança e rastreabilidade..."):
                prompt = f"""
                Atue como PMO Sênior SAP. Analise a transcrição abaixo e gere uma ATA EXECUTIVA.
                Regras: Use IDs (T1, A1, R1) para rastreabilidade total entre temas, ações e riscos.
                Use Tabelas Markdown para os itens 1, 2 e 3. Use tabela para os proximos passos.
                
                Transcrição: {transcricao}
                """
                response = model.generate_content(prompt)
                ata_final = response.text

            # --- FASE 3: RESULTADOS ---
            st.success("✅ Ata gerada com sucesso!")
            st.markdown(ata_final)
            
            btn_download = st.download_button(
                label="⬇️ Baixar em Word (.docx)",
                data=criar_docx(ata_final),
                file_name="Ata_Executiva.docx"
            )
            
            os.remove("temp.mp3") # Limpa o arquivo temporário

        except Exception as e:
            st.error(f"Erro no processamento: {e}")